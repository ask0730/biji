"""
使用百度 OCR 扫描 PDF 文件，并将结果输出到 Word 文档
"""
import os
import sys
import json
import logging
from pathlib import Path
from datetime import datetime

# 设置 Windows 控制台编码为 UTF-8
if sys.platform == 'win32':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except:
        pass

# 检查必要的库
try:
    from aip import AipOcr
    BAIDU_AIP_AVAILABLE = True
except ImportError:
    BAIDU_AIP_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


def init_logger():
    """初始化日志"""
    log_dir = os.path.join(os.path.dirname(__file__), "logs")
    os.makedirs(log_dir, exist_ok=True)
    log_path = os.path.join(log_dir, "baidu_ocr.log")

    file_handler = logging.FileHandler(log_path, encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
    file_handler.setFormatter(formatter)
    
    logger = logging.getLogger("baidu_ocr")
    logger.setLevel(logging.INFO)
    logger.addHandler(file_handler)
    
    # 同时输出到控制台
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(formatter)
    logger.addHandler(console_handler)
    
    return logger


logger = init_logger()


def check_dependencies():
    """检查依赖库"""
    missing = []
    
    if not BAIDU_AIP_AVAILABLE:
        missing.append("baidu-aip (pip install baidu-aip)")
    
    if not PYMUPDF_AVAILABLE:
        missing.append("PyMuPDF (pip install PyMuPDF)")
    
    if not DOCX_AVAILABLE:
        missing.append("python-docx (pip install python-docx)")
    
    if not PIL_AVAILABLE:
        missing.append("Pillow (pip install Pillow)")
    
    if missing:
        print("错误: 缺少以下依赖库:")
        for lib in missing:
            print(f"  - {lib}")
        print("\n请先安装这些库后再运行脚本")
        return False
    
    return True


def load_config_from_file(script_dir):
    """从配置文件加载 API 密钥"""
    config_file = os.path.join(script_dir, 'baidu_ocr_config.json')
    if os.path.exists(config_file):
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
                # 去除可能的空格
                app_id = str(config.get('APP_ID', '')).strip()
                api_key = str(config.get('API_KEY', '')).strip()
                secret_key = str(config.get('SECRET_KEY', '')).strip()
                return (app_id, api_key, secret_key)
        except Exception as e:
            logger.warning(f"读取配置文件失败: {e}")
    return ('', '', '')


def init_baidu_ocr():
    """初始化百度 OCR 客户端"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 优先级: 环境变量 > 配置文件 > 用户输入
    APP_ID = os.getenv('BAIDU_OCR_APP_ID', '').strip()
    API_KEY = os.getenv('BAIDU_OCR_API_KEY', '').strip()
    SECRET_KEY = os.getenv('BAIDU_OCR_SECRET_KEY', '').strip()
    
    # 如果环境变量未设置，尝试从配置文件读取
    if not APP_ID or not API_KEY or not SECRET_KEY:
        file_app_id, file_api_key, file_secret_key = load_config_from_file(script_dir)
        if not APP_ID:
            APP_ID = file_app_id
        if not API_KEY:
            API_KEY = file_api_key
        if not SECRET_KEY:
            SECRET_KEY = file_secret_key
    
    # 如果仍未设置，提示用户输入
    if not APP_ID or not API_KEY or not SECRET_KEY:
        print("\n" + "="*60)
        print("需要配置百度 OCR API 密钥")
        print("="*60)
        print("您可以通过以下方式配置:")
        print("1. 设置环境变量:")
        print("   - BAIDU_OCR_APP_ID")
        print("   - BAIDU_OCR_API_KEY")
        print("   - BAIDU_OCR_SECRET_KEY")
        print("\n2. 创建配置文件 baidu_ocr_config.json:")
        print('   {"APP_ID": "your_app_id", "API_KEY": "your_api_key", "SECRET_KEY": "your_secret_key"}')
        print("\n3. 或直接在提示时输入")
        print("\n获取 API 密钥: https://console.bce.baidu.com/ai/#/ai/ocr/overview/index")
        print("="*60)
        
        if not APP_ID:
            APP_ID = input("请输入 APP_ID: ").strip()
        if not API_KEY:
            API_KEY = input("请输入 API_KEY: ").strip()
        if not SECRET_KEY:
            SECRET_KEY = input("请输入 SECRET_KEY: ").strip()
    
    if not APP_ID or not API_KEY or not SECRET_KEY:
        raise ValueError("百度 OCR API 密钥未配置")
    
    # 记录密钥信息（不记录完整密钥，只记录长度和前后几位）
    logger.info(f"使用 APP_ID: {APP_ID}")
    logger.info(f"使用 API_KEY: {API_KEY[:8]}...{API_KEY[-4:] if len(API_KEY) > 12 else '***'}")
    logger.info(f"使用 SECRET_KEY: {SECRET_KEY[:8]}...{SECRET_KEY[-4:] if len(SECRET_KEY) > 12 else '***'}")
    
    client = AipOcr(APP_ID, API_KEY, SECRET_KEY)
    logger.info("百度 OCR 客户端初始化成功")
    return client


def pdf_to_images(pdf_path, dpi=200):
    """将 PDF 转换为图片列表（使用 PyMuPDF）"""
    logger.info(f"开始将 PDF 转换为图片: {pdf_path}")
    try:
        doc = fitz.open(pdf_path)
        images = []
        
        # 计算缩放比例（dpi 转换为缩放因子）
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            # 将页面渲染为图片
            pix = page.get_pixmap(matrix=mat)
            # 转换为 PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            images.append(img)
        
        doc.close()
        logger.info(f"成功转换 {len(images)} 页")
        return images
    except Exception as e:
        logger.error(f"PDF 转换失败: {e}")
        raise


def image_to_bytes(image):
    """将 PIL Image 转换为字节流"""
    import io
    img_byte_arr = io.BytesIO()
    image.save(img_byte_arr, format='PNG')
    img_byte_arr = img_byte_arr.getvalue()
    return img_byte_arr


def ocr_image(client, image_bytes):
    """使用百度 OCR 识别图片"""
    try:
        result = client.basicGeneral(image_bytes)
        
        if 'error_code' in result:
            error_code = result.get('error_code', '')
            error_msg = result.get('error_msg', '未知错误')
            logger.error(f"OCR 识别失败: {error_code} - {error_msg}")
            
            # 如果是认证错误，给出更详细的提示
            if '110' in str(error_code) or 'IAM' in error_msg or '认证' in error_msg:
                logger.error("API 认证失败，请检查 APP_ID、API_KEY 和 SECRET_KEY 是否正确")
            
            return None
        
        return result
    except Exception as e:
        logger.error(f"OCR 识别异常: {e}")
        return None


def extract_text_from_ocr_result(result):
    """从 OCR 结果中提取文本"""
    if not result or 'words_result' not in result:
        return ""
    
    texts = []
    for item in result['words_result']:
        if 'words' in item:
            texts.append(item['words'])
    
    return '\n'.join(texts)


def create_word_document(ocr_results, output_path):
    """创建 Word 文档并写入 OCR 结果"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('PDF OCR 识别结果', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加时间戳
    time_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    time_para = doc.add_paragraph(f'识别时间: {time_str}')
    time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # 遍历每一页的结果
    for page_num, result in enumerate(ocr_results, 1):
        # 添加页码标题
        page_heading = doc.add_heading(f'第 {page_num} 页', level=1)
        
        if result is None:
            doc.add_paragraph('识别失败，请检查图片质量或 API 配置')
            doc.add_paragraph()  # 空行
            continue
        
        # 提取文本
        text = extract_text_from_ocr_result(result)
        
        if text:
            # 添加文本内容
            para = doc.add_paragraph(text)
            # 设置字体
            for run in para.runs:
                run.font.name = '宋体'
                run.font.size = Pt(12)
        else:
            doc.add_paragraph('未识别到文本内容')
        
        # 添加详细信息（可选）
        if 'words_result_num' in result:
            word_count = result['words_result_num']
            info_para = doc.add_paragraph(f'识别到 {word_count} 个文本块')
            info_para.runs[0].font.size = Pt(10)
            info_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # 空行分隔每一页
        doc.add_page_break()  # 分页符
    
    # 保存文档
    doc.save(output_path)
    logger.info(f"Word 文档已保存: {output_path}")


def main():
    """主函数"""
    # 检查依赖
    if not check_dependencies():
        sys.exit(1)
    
    try:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        
        # 查找 PDF 文件
        pdf_file = os.path.join(script_dir, "职称参评-毛雅君20251114V1.0(1).pdf")
        
        if not os.path.exists(pdf_file):
            print(f"错误: 未找到文件 {pdf_file}")
            logger.error(f"未找到文件: {pdf_file}")
            sys.exit(1)
        
        print(f"\n找到 PDF 文件: {os.path.basename(pdf_file)}")
        logger.info(f"开始处理文件: {pdf_file}")
        
        # 初始化百度 OCR
        print("\n正在初始化百度 OCR...")
        client = init_baidu_ocr()
        
        # 将 PDF 转换为图片
        print("\n正在将 PDF 转换为图片...")
        images = pdf_to_images(pdf_file)
        print(f"共 {len(images)} 页")
        
        # OCR 识别每一页
        print("\n开始 OCR 识别...")
        ocr_results = []
        
        for i, image in enumerate(images, 1):
            print(f"正在识别第 {i}/{len(images)} 页...")
            logger.info(f"正在识别第 {i} 页")
            
            # 转换为字节流
            image_bytes = image_to_bytes(image)
            
            # OCR 识别
            result = ocr_image(client, image_bytes)
            ocr_results.append(result)
            
            if result:
                word_count = result.get('words_result_num', 0)
                print(f"  识别到 {word_count} 个文本块")
            else:
                print(f"  识别失败")
        
        # 生成输出文件名
        pdf_name = Path(pdf_file).stem
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_file = os.path.join(script_dir, f"{pdf_name}_OCR结果_{timestamp}.docx")
        
        # 创建 Word 文档
        print(f"\n正在生成 Word 文档...")
        create_word_document(ocr_results, output_file)
        
        try:
            print(f"\n完成！结果已保存到: {os.path.basename(output_file)}")
        except UnicodeEncodeError:
            print(f"\n完成！结果已保存到: {output_file}")
        logger.info(f"处理完成，输出文件: {output_file}")
        
    except KeyboardInterrupt:
        print("\n\n用户中断操作")
        logger.info("用户中断操作")
        sys.exit(0)
    except Exception as e:
        print(f"\n错误: {e}")
        logger.error(f"处理失败: {e}", exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()


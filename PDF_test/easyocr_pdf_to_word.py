"""
使用 EasyOCR 扫描 PDF 文件，并将结果输出到 Word 文档
EasyOCR 是免费的，中文识别效果好
"""
import os
import sys
import logging
from pathlib import Path
from datetime import datetime

# 设置 Windows 控制台编码为 UTF-8
if sys.platform == 'win32':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except:
        pass

# 检查必要的库
try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image, ImageEnhance, ImageFilter
    import numpy as np
    import cv2
    PIL_AVAILABLE = True
    CV2_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    CV2_AVAILABLE = False


def init_logger():
    """初始化日志"""
    log_dir = os.path.join(os.path.dirname(__file__), "logs")
    os.makedirs(log_dir, exist_ok=True)
    log_path = os.path.join(log_dir, "easyocr.log")

    file_handler = logging.FileHandler(log_path, encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
    file_handler.setFormatter(formatter)
    
    logger = logging.getLogger("easyocr")
    logger.setLevel(logging.INFO)
    logger.addHandler(file_handler)
    
    # 同时输出到控制台
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(formatter)
    logger.addHandler(console_handler)
    
    return logger


logger = init_logger()


def check_dependencies():
    """检查依赖库"""
    missing = []
    
    if not EASYOCR_AVAILABLE:
        missing.append("easyocr (pip install easyocr)")
    
    if not PYMUPDF_AVAILABLE:
        missing.append("PyMuPDF (pip install PyMuPDF)")
    
    if not DOCX_AVAILABLE:
        missing.append("python-docx (pip install python-docx)")
    
    if not PIL_AVAILABLE:
        missing.append("Pillow (pip install Pillow)")
    
    if not CV2_AVAILABLE:
        missing.append("opencv-python (pip install opencv-python)")
    
    if missing:
        print("错误: 缺少以下依赖库:")
        for lib in missing:
            print(f"  - {lib}")
        print("\n请先安装这些库后再运行脚本")
        print("\n安装命令:")
        print("  pip install easyocr PyMuPDF python-docx Pillow numpy")
        print("\n注意: EasyOCR 首次运行会自动下载模型，需要一些时间")
        return False
    
    return True


def init_easyocr():
    """初始化 EasyOCR"""
    logger.info("正在初始化 EasyOCR（首次运行会下载模型，请耐心等待）...")
    print("正在初始化 EasyOCR...")
    print("提示: 首次运行会自动下载中文和英文模型文件，可能需要几分钟，请耐心等待...")
    
    try:
        # 使用中文和英文，启用 GPU（如果可用）
        reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)
        logger.info("EasyOCR 初始化成功")
        print("EasyOCR 初始化成功！")
        return reader
    except Exception as e:
        logger.error(f"EasyOCR 初始化失败: {e}")
        raise


def pdf_to_images(pdf_path, dpi=300):
    """将 PDF 转换为图片列表（使用 PyMuPDF），提高 DPI 以提高识别准确率"""
    logger.info(f"开始将 PDF 转换为图片: {pdf_path} (DPI: {dpi})")
    try:
        doc = fitz.open(pdf_path)
        images = []
        
        # 计算缩放比例（dpi 转换为缩放因子）
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            # 将页面渲染为图片
            pix = page.get_pixmap(matrix=mat)
            # 转换为 PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            images.append(img)
        
        doc.close()
        logger.info(f"成功转换 {len(images)} 页")
        return images
    except Exception as e:
        logger.error(f"PDF 转换失败: {e}")
        raise


def preprocess_image(image):
    """图片预处理以提高 OCR 识别准确率"""
    try:
        # 转换为 numpy 数组
        img_array = np.array(image)
        
        # 转换为 OpenCV 格式 (BGR)
        if len(img_array.shape) == 3:
            img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
        else:
            img_cv = img_array
        
        # 1. 转为灰度图
        if len(img_cv.shape) == 3:
            gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
        else:
            gray = img_cv
        
        # 2. 降噪
        denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
        
        # 3. 增强对比度（使用 CLAHE - 对比度受限的自适应直方图均衡化）
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(denoised)
        
        # 4. 锐化
        kernel = np.array([[-1, -1, -1],
                          [-1,  9, -1],
                          [-1, -1, -1]])
        sharpened = cv2.filter2D(enhanced, -1, kernel)
        
        # 5. 二值化（自适应阈值）
        _, binary = cv2.threshold(sharpened, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # 转换回 RGB 格式（EasyOCR 需要）
        if len(binary.shape) == 2:
            # 灰度图转 RGB
            processed = cv2.cvtColor(binary, cv2.COLOR_GRAY2RGB)
        else:
            processed = binary
        
        return processed
    
    except Exception as e:
        logger.warning(f"图片预处理失败，使用原图: {e}")
        return np.array(image)


def image_to_array(image, preprocess=True):
    """将 PIL Image 转换为 numpy 数组（EasyOCR 需要的格式）"""
    if preprocess and CV2_AVAILABLE:
        # 使用预处理后的图片
        processed = preprocess_image(image)
        return processed
    else:
        # 直接转换
        return np.array(image)


def ocr_image(reader, image_array, min_confidence=0.1):
    """使用 EasyOCR 识别图片，极低置信度阈值以保留所有内容（特别是表格）"""
    try:
        # 设置参数以识别更多内容，特别是表格中的小字
        result = reader.readtext(
            image_array,
            detail=1,  # 返回详细信息（包含坐标）
            paragraph=False,  # 不自动合并段落，保持原始位置
            width_ths=0.3,  # 进一步降低文本宽度阈值
            height_ths=0.3,  # 进一步降低文本高度阈值
            allowlist=None,  # 允许所有字符
        )
        
        # 极低置信度阈值以保留所有内容（表格内容可能置信度较低）
        filtered_result = []
        for detection in result:
            if len(detection) >= 3:
                confidence = detection[2]
                text = detection[1] if len(detection) >= 2 else ''
                # 只要文本不为空，就保留（置信度阈值很低）
                if text and text.strip() and confidence >= min_confidence:
                    filtered_result.append(detection)
            elif len(detection) >= 2:
                # 如果没有置信度信息，也保留（可能是表格中的内容）
                text = detection[1]
                if text and text.strip():
                    filtered_result.append(detection)
        
        return filtered_result if filtered_result else None
    except Exception as e:
        logger.error(f"OCR 识别异常: {e}")
        return None


def extract_text_from_ocr_result(result, min_confidence=0.1):
    """从 EasyOCR 结果中提取文本，极低阈值以保留所有内容"""
    if not result:
        return ""
    
    texts = []
    for detection in result:
        if len(detection) >= 2:
            text = detection[1]  # 提取文本内容
            confidence = detection[2] if len(detection) >= 3 else 1.0  # 置信度
            
            # 极低阈值以保留所有内容（特别是表格中的内容）
            if text and text.strip() and confidence >= min_confidence:
                texts.append(text)
    
    return '\n'.join(texts)


def is_likely_table(result, min_rows=3, min_cols=2):
    """判断识别结果是否可能是表格"""
    if not result or len(result) < min_rows * min_cols:
        return False
    
    # 检测表格结构
    table_data = detect_table_structure(result, row_tolerance=25, col_tolerance=100)
    if table_data and len(table_data) >= min_rows:
        # 检查是否有足够的列
        max_cols = max([len(row) for row in table_data if row])
        return max_cols >= min_cols
    
    return False


def detect_table_structure(result, row_tolerance=20, col_tolerance=80):
    """检测表格结构，根据文本位置坐标分组为行和列"""
    if not result:
        return []
    
    # 提取所有文本及其坐标
    items = []
    for detection in result:
        if len(detection) >= 2:
            bbox = detection[0]  # 边界框坐标 [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
            text = detection[1]
            confidence = detection[2] if len(detection) >= 3 else 1.0
            
            # 计算文本的中心点和边界
            if bbox and len(bbox) >= 4:
                x_coords = [point[0] for point in bbox]
                y_coords = [point[1] for point in bbox]
                center_x = sum(x_coords) / len(x_coords)
                center_y = sum(y_coords) / len(y_coords)
                min_x = min(x_coords)
                max_x = max(x_coords)
                min_y = min(y_coords)
                max_y = max(y_coords)
                
                items.append({
                    'text': text,
                    'x': center_x,
                    'y': center_y,
                    'min_x': min_x,
                    'max_x': max_x,
                    'min_y': min_y,
                    'max_y': max_y,
                    'confidence': confidence
                })
    
    if not items:
        return []
    
    # 按Y坐标分组为行（使用容差）
    rows = []
    sorted_items = sorted(items, key=lambda item: item['y'])
    
    current_row = []
    current_y = None
    
    for item in sorted_items:
        if current_y is None or abs(item['y'] - current_y) <= row_tolerance:
            # 同一行
            current_row.append(item)
            if current_y is None:
                current_y = item['y']
            else:
                current_y = (current_y + item['y']) / 2  # 更新平均Y坐标
        else:
            # 新行
            if current_row:
                rows.append(current_row)
            current_row = [item]
            current_y = item['y']
    
    if current_row:
        rows.append(current_row)
    
    # 对每一行内的项目按X坐标排序
    for row in rows:
        row.sort(key=lambda item: item['x'])
    
    # 检测列结构（找出所有可能的列位置）
    all_x_positions = sorted(set([item['x'] for row in rows for item in row]))
    
    # 合并相近的X位置作为列
    columns = []
    for x in all_x_positions:
        # 检查是否与已有列相近
        found = False
        for col_x in columns:
            if abs(x - col_x) <= col_tolerance:
                found = True
                break
        if not found:
            columns.append(x)
    
    columns.sort()
    
    # 构建表格数据
    table_data = []
    for row in rows:
        row_data = [''] * len(columns)
        for item in row:
            # 找到最接近的列
            best_col = 0
            min_dist = float('inf')
            for i, col_x in enumerate(columns):
                dist = abs(item['x'] - col_x)
                if dist < min_dist:
                    min_dist = dist
                    best_col = i
            
            # 如果文本已存在，合并（用空格分隔）
            if row_data[best_col]:
                row_data[best_col] += ' ' + item['text']
            else:
                row_data[best_col] = item['text']
        
        table_data.append(row_data)
    
    return table_data


def create_word_document(ocr_results, output_path):
    """创建 Word 文档并写入 OCR 结果，尝试识别表格结构"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('PDF OCR 识别结果（EasyOCR - 表格模式）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加时间戳
    time_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    time_para = doc.add_paragraph(f'识别时间: {time_str}')
    time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # 遍历每一页的结果
    for page_num, result in enumerate(ocr_results, 1):
        # 添加页码标题
        page_heading = doc.add_heading(f'第 {page_num} 页', level=1)
        
        if result is None or not result:
            doc.add_paragraph('识别失败或未识别到内容')
            doc.add_paragraph()  # 空行
            continue
        
        # 尝试检测表格结构（使用多种容差策略）
        # 先尝试较小的容差（适合紧凑表格）
        table_data = detect_table_structure(result, row_tolerance=25, col_tolerance=100)
        
        # 如果检测不到表格或表格太小，尝试中等容差
        if not table_data or (table_data and len(table_data) < 2):
            table_data = detect_table_structure(result, row_tolerance=35, col_tolerance=130)
        
        # 如果还是检测不到，尝试更大的容差（适合稀疏表格）
        if not table_data or (table_data and len(table_data) < 2):
            table_data = detect_table_structure(result, row_tolerance=50, col_tolerance=180)
        
        if table_data and len(table_data) > 0:
            # 创建表格
            num_cols = max(len(row) for row in table_data) if table_data else 1
            num_rows = len(table_data)
            
            if num_cols > 1 and num_rows >= 2:
                # 创建Word表格
                logger.info(f"第 {page_num} 页创建表格: {num_rows} 行 x {num_cols} 列")
                print(f"  创建表格: {num_rows} 行 x {num_cols} 列")
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'
                
                # 填充表格数据
                for row_idx, row_data in enumerate(table_data):
                    for col_idx in range(num_cols):
                        cell = table.rows[row_idx].cells[col_idx]
                        if col_idx < len(row_data):
                            cell_text = str(row_data[col_idx]).strip() if row_data[col_idx] else ''
                            cell.text = cell_text
                        else:
                            cell.text = ''
                        
                        # 设置字体
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = '宋体'
                                run.font.size = Pt(9)
                
                # 添加表格说明
                info_para = doc.add_paragraph(f'表格: {num_rows} 行 x {num_cols} 列')
                info_para.runs[0].font.size = Pt(9)
                info_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            else:
                # 如果检测不到表格结构，使用文本形式
                text = extract_text_from_ocr_result(result, min_confidence=0.1)
                if text:
                    para = doc.add_paragraph(text)
                    for run in para.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(12)
        else:
            # 如果检测不到表格结构，使用文本形式
            text = extract_text_from_ocr_result(result, min_confidence=0.1)
            if text:
                para = doc.add_paragraph(text)
                for run in para.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(12)
            else:
                doc.add_paragraph('未识别到文本内容')
        
        # 添加统计信息
        if result:
            item_count = len([d for d in result if d and len(d) >= 2])
            info_para = doc.add_paragraph(f'识别到 {item_count} 个文本块')
            info_para.runs[0].font.size = Pt(10)
            info_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # 空行分隔每一页
        doc.add_page_break()  # 分页符
    
    # 保存文档
    doc.save(output_path)
    logger.info(f"Word 文档已保存: {output_path}")


def main():
    """主函数"""
    # 检查依赖
    if not check_dependencies():
        sys.exit(1)
    
    try:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        
        # 查找 PDF 文件
        pdf_file = os.path.join(script_dir, "职称参评-毛雅君20251114V1.0(1).pdf")
        
        if not os.path.exists(pdf_file):
            print(f"错误: 未找到文件 {pdf_file}")
            logger.error(f"未找到文件: {pdf_file}")
            sys.exit(1)
        
        print(f"\n找到 PDF 文件: {os.path.basename(pdf_file)}")
        logger.info(f"开始处理文件: {pdf_file}")
        
        # 初始化 EasyOCR
        print("\n正在初始化 EasyOCR...")
        reader = init_easyocr()
        
        # 将 PDF 转换为图片
        print("\n正在将 PDF 转换为图片...")
        images = pdf_to_images(pdf_file)
        print(f"共 {len(images)} 页")
        
        # OCR 识别每一页
        print("\n开始 OCR 识别...")
        ocr_results = []
        
        for i, image in enumerate(images, 1):
            print(f"正在识别第 {i}/{len(images)} 页...")
            logger.info(f"正在识别第 {i} 页")
            
            # 转换为 numpy 数组（启用预处理）
            image_array = image_to_array(image, preprocess=True)
            
            # OCR 识别（极低置信度阈值以保留所有内容，特别是表格）
            result = ocr_image(reader, image_array, min_confidence=0.1)
            ocr_results.append(result)
            
            if result:
                line_count = len([d for d in result if d and len(d) >= 2])
                print(f"  识别到 {line_count} 行文本")
            else:
                print(f"  识别失败或未识别到内容")
        
        # 生成输出文件名
        pdf_name = Path(pdf_file).stem
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_file = os.path.join(script_dir, f"{pdf_name}_OCR结果_{timestamp}.docx")
        
        # 创建 Word 文档
        print(f"\n正在生成 Word 文档...")
        create_word_document(ocr_results, output_file)
        
        try:
            print(f"\n完成！结果已保存到: {os.path.basename(output_file)}")
        except UnicodeEncodeError:
            print(f"\n完成！结果已保存到: {output_file}")
        logger.info(f"处理完成，输出文件: {output_file}")
        
    except KeyboardInterrupt:
        print("\n\n用户中断操作")
        logger.info("用户中断操作")
        sys.exit(0)
    except Exception as e:
        print(f"\n错误: {e}")
        logger.error(f"处理失败: {e}", exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()



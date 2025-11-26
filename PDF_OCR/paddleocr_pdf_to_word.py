"""
使用 PaddleOCR 扫描 PDF 文件，并将结果输出到 Word 文档
PaddleOCR 是免费的，中文识别效果好
"""
import os
import sys
import logging
from pathlib import Path
from datetime import datetime

# 设置 Windows 控制台编码为 UTF-8
if sys.platform == 'win32':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except:
        pass

# 检查必要的库
try:
    from paddleocr import PaddleOCR
    PADDLEOCR_AVAILABLE = True
except ImportError:
    PADDLEOCR_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image
    import numpy as np
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


def init_logger():
    """初始化日志"""
    log_dir = os.path.join(os.path.dirname(__file__), "logs")
    os.makedirs(log_dir, exist_ok=True)
    log_path = os.path.join(log_dir, "paddleocr.log")

    file_handler = logging.FileHandler(log_path, encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
    file_handler.setFormatter(formatter)
    
    logger = logging.getLogger("paddleocr")
    logger.setLevel(logging.INFO)
    logger.addHandler(file_handler)
    
    # 同时输出到控制台
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(formatter)
    logger.addHandler(console_handler)
    
    return logger


logger = init_logger()


def check_dependencies():
    """检查依赖库"""
    missing = []
    
    if not PADDLEOCR_AVAILABLE:
        missing.append("paddleocr (pip install paddleocr)")
    
    if not PYMUPDF_AVAILABLE:
        missing.append("PyMuPDF (pip install PyMuPDF)")
    
    if not DOCX_AVAILABLE:
        missing.append("python-docx (pip install python-docx)")
    
    if not PIL_AVAILABLE:
        missing.append("Pillow (pip install Pillow)")
    
    if missing:
        print("错误: 缺少以下依赖库:")
        for lib in missing:
            print(f"  - {lib}")
        print("\n请先安装这些库后再运行脚本")
        print("\n安装命令:")
        print("  pip install paddleocr PyMuPDF python-docx Pillow")
        print("\n注意: PaddleOCR 首次运行会自动下载模型，需要一些时间")
        return False
    
    return True


def init_paddleocr():
    """初始化 PaddleOCR"""
    logger.info("正在初始化 PaddleOCR（首次运行会下载模型，请耐心等待）...")
    print("正在初始化 PaddleOCR...")
    print("提示: 首次运行会自动下载模型文件，可能需要几分钟，请耐心等待...")
    
    try:
        # 使用中文模型，启用 GPU（如果可用）
        ocr = PaddleOCR(use_angle_cls=True, lang='ch', use_gpu=False)
        logger.info("PaddleOCR 初始化成功")
        print("PaddleOCR 初始化成功！")
        return ocr
    except Exception as e:
        logger.error(f"PaddleOCR 初始化失败: {e}")
        raise


def pdf_to_images(pdf_path, dpi=200):
    """将 PDF 转换为图片列表（使用 PyMuPDF）"""
    logger.info(f"开始将 PDF 转换为图片: {pdf_path}")
    try:
        doc = fitz.open(pdf_path)
        images = []
        
        # 计算缩放比例（dpi 转换为缩放因子）
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            # 将页面渲染为图片
            pix = page.get_pixmap(matrix=mat)
            # 转换为 PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            images.append(img)
        
        doc.close()
        logger.info(f"成功转换 {len(images)} 页")
        return images
    except Exception as e:
        logger.error(f"PDF 转换失败: {e}")
        raise


def image_to_array(image):
    """将 PIL Image 转换为 numpy 数组（PaddleOCR 需要的格式）"""
    return np.array(image)


def ocr_image(ocr, image_array):
    """使用 PaddleOCR 识别图片"""
    try:
        result = ocr.ocr(image_array, cls=True)
        return result
    except Exception as e:
        logger.error(f"OCR 识别异常: {e}")
        return None


def extract_text_from_ocr_result(result):
    """从 PaddleOCR 结果中提取文本"""
    if not result or not result[0]:
        return ""
    
    texts = []
    for line in result[0]:
        if line and len(line) >= 2:
            text = line[1][0]  # 提取文本内容
            confidence = line[1][1]  # 置信度
            if text:
                texts.append(text)
    
    return '\n'.join(texts)


def create_word_document(ocr_results, output_path):
    """创建 Word 文档并写入 OCR 结果"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('PDF OCR 识别结果（PaddleOCR）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加时间戳
    time_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    time_para = doc.add_paragraph(f'识别时间: {time_str}')
    time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # 遍历每一页的结果
    for page_num, result in enumerate(ocr_results, 1):
        # 添加页码标题
        page_heading = doc.add_heading(f'第 {page_num} 页', level=1)
        
        if result is None or not result or not result[0]:
            doc.add_paragraph('识别失败或未识别到内容')
            doc.add_paragraph()  # 空行
            continue
        
        # 提取文本
        text = extract_text_from_ocr_result(result)
        
        if text:
            # 添加文本内容
            para = doc.add_paragraph(text)
            # 设置字体
            for run in para.runs:
                run.font.name = '宋体'
                run.font.size = Pt(12)
        else:
            doc.add_paragraph('未识别到文本内容')
        
        # 添加统计信息
        if result and result[0]:
            line_count = len([line for line in result[0] if line and len(line) >= 2])
            info_para = doc.add_paragraph(f'识别到 {line_count} 行文本')
            info_para.runs[0].font.size = Pt(10)
            info_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # 空行分隔每一页
        doc.add_page_break()  # 分页符
    
    # 保存文档
    doc.save(output_path)
    logger.info(f"Word 文档已保存: {output_path}")


def main():
    """主函数"""
    # 检查依赖
    if not check_dependencies():
        sys.exit(1)
    
    try:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        
        # 查找 PDF 文件
        pdf_file = os.path.join(script_dir, "职称参评-毛雅君20251114V1.0(1).pdf")
        
        if not os.path.exists(pdf_file):
            print(f"错误: 未找到文件 {pdf_file}")
            logger.error(f"未找到文件: {pdf_file}")
            sys.exit(1)
        
        print(f"\n找到 PDF 文件: {os.path.basename(pdf_file)}")
        logger.info(f"开始处理文件: {pdf_file}")
        
        # 初始化 PaddleOCR
        print("\n正在初始化 PaddleOCR...")
        ocr = init_paddleocr()
        
        # 将 PDF 转换为图片
        print("\n正在将 PDF 转换为图片...")
        images = pdf_to_images(pdf_file)
        print(f"共 {len(images)} 页")
        
        # OCR 识别每一页
        print("\n开始 OCR 识别...")
        ocr_results = []
        
        for i, image in enumerate(images, 1):
            print(f"正在识别第 {i}/{len(images)} 页...")
            logger.info(f"正在识别第 {i} 页")
            
            # 转换为 numpy 数组
            image_array = image_to_array(image)
            
            # OCR 识别
            result = ocr_image(ocr, image_array)
            ocr_results.append(result)
            
            if result and result[0]:
                line_count = len([line for line in result[0] if line and len(line) >= 2])
                print(f"  识别到 {line_count} 行文本")
            else:
                print(f"  识别失败或未识别到内容")
        
        # 生成输出文件名
        pdf_name = Path(pdf_file).stem
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_file = os.path.join(script_dir, f"{pdf_name}_OCR结果_{timestamp}.docx")
        
        # 创建 Word 文档
        print(f"\n正在生成 Word 文档...")
        create_word_document(ocr_results, output_file)
        
        try:
            print(f"\n完成！结果已保存到: {os.path.basename(output_file)}")
        except UnicodeEncodeError:
            print(f"\n完成！结果已保存到: {output_file}")
        logger.info(f"处理完成，输出文件: {output_file}")
        
    except KeyboardInterrupt:
        print("\n\n用户中断操作")
        logger.info("用户中断操作")
        sys.exit(0)
    except Exception as e:
        print(f"\n错误: {e}")
        logger.error(f"处理失败: {e}", exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()




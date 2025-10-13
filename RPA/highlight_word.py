#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档直接高亮工具
直接在Word文档中高亮显示邮箱和年份
"""

import re
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from pathlib import Path

def highlight_word_document(file_path):
    """在Word文档中高亮邮箱和年份"""
    email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
    year_pattern = re.compile(r'\b(19|20)\d{2}\b')
    
    print(f"处理文件: {file_path.name}")
    
    try:
        # 打开文档
        doc = Document(file_path)
        
        # 创建备份
        backup_path = file_path.with_suffix('.docx.backup')
        if not backup_path.exists():
            doc.save(backup_path)
            print(f"  已创建备份: {backup_path.name}")
        
        email_count = 0
        year_count = 0
        
        # 处理段落
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if not text:
                continue
                
            # 查找邮箱
            email_matches = list(email_pattern.finditer(text))
            # 查找年份
            year_matches = list(year_pattern.finditer(text))
            
            if email_matches or year_matches:
                # 清除段落
                paragraph.clear()
                
                last_end = 0
                
                # 合并所有匹配项并按位置排序
                all_matches = []
                for match in email_matches:
                    all_matches.append((match.start(), match.end(), 'email'))
                for match in year_matches:
                    all_matches.append((match.start(), match.end(), 'year'))
                
                all_matches.sort(key=lambda x: x[0])
                
                # 重新构建段落
                for start, end, match_type in all_matches:
                    # 添加匹配前的文本
                    if start > last_end:
                        paragraph.add_run(text[last_end:start])
                    
                    # 添加高亮的文本
                    run = paragraph.add_run(text[start:end])
                    if match_type == 'email':
                        run.font.highlight_color = WD_COLOR_INDEX.GREEN
                        email_count += 1
                    else:  # year
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        year_count += 1
                    
                    last_end = end
                
                # 添加最后剩余的文本
                if last_end < len(text):
                    paragraph.add_run(text[last_end:])
        
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if not text:
                            continue
                            
                        # 查找邮箱和年份
                        email_matches = list(email_pattern.finditer(text))
                        year_matches = list(year_pattern.finditer(text))
                        
                        if email_matches or year_matches:
                            # 清除段落
                            paragraph.clear()
                            
                            last_end = 0
                            
                            # 合并所有匹配项并按位置排序
                            all_matches = []
                            for match in email_matches:
                                all_matches.append((match.start(), match.end(), 'email'))
                            for match in year_matches:
                                all_matches.append((match.start(), match.end(), 'year'))
                            
                            all_matches.sort(key=lambda x: x[0])
                            
                            # 重新构建段落
                            for start, end, match_type in all_matches:
                                # 添加匹配前的文本
                                if start > last_end:
                                    paragraph.add_run(text[last_end:start])
                                
                                # 添加高亮的文本
                                run = paragraph.add_run(text[start:end])
                                if match_type == 'email':
                                    run.font.highlight_color = WD_COLOR_INDEX.GREEN
                                    email_count += 1
                                else:  # year
                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                    year_count += 1
                                
                                last_end = end
                            
                            # 添加最后剩余的文本
                            if last_end < len(text):
                                paragraph.add_run(text[last_end:])
        
        # 保存文档
        doc.save(file_path)
        
        print(f"  ✅ 完成高亮处理")
        print(f"  📧 邮箱数量: {email_count}")
        print(f"  📅 年份数量: {year_count}")
        
        return email_count, year_count
        
    except Exception as e:
        print(f"  ❌ 处理失败: {e}")
        return 0, 0

def main():
    """主函数"""
    directory_path = Path(r"D:\Desktop\图书馆\项目申报")
    
    print("=" * 60)
    print("Word文档直接高亮工具")
    print("=" * 60)
    print(f"目标目录: {directory_path}")
    print()
    
    if not directory_path.exists():
        print(f"❌ 目录不存在: {directory_path}")
        return
    
    # 查找Word文件
    word_files = list(directory_path.glob("*.docx"))
    if not word_files:
        print("❌ 未找到Word文件")
        return
    
    print(f"📁 找到 {len(word_files)} 个Word文件")
    print()
    
    # 确认操作
    print("⚠️  警告: 将直接修改Word文档!")
    print("🎨 高亮效果:")
    print("  🟢 绿色高亮: 邮箱地址")
    print("  🟡 黄色高亮: 年份信息")
    print()
    
    try:
        confirm = input("确认继续? (y/n): ").lower().strip()
        if confirm not in ['y', 'yes', '是']:
            print("操作已取消")
            return
    except:
        print("操作已取消")
        return
    
    print()
    print("开始处理...")
    print()
    
    total_emails = 0
    total_years = 0
    processed_files = 0
    
    # 处理每个文件
    for file_path in word_files:
        email_count, year_count = highlight_word_document(file_path)
        total_emails += email_count
        total_years += year_count
        if email_count > 0 or year_count > 0:
            processed_files += 1
        print()
    
    # 显示结果
    print("=" * 60)
    print("处理完成!")
    print("=" * 60)
    print(f"📊 处理文件数: {len(word_files)}")
    print(f"✅ 成功处理: {processed_files}")
    print(f"📧 总邮箱数量: {total_emails}")
    print(f"📅 总年份数量: {total_years}")
    print()
    print("💡 提示:")
    print("  - 原文件已备份为 .backup 文件")
    print("  - 可以直接在Word中查看高亮效果")
    print("  - 如需恢复，删除高亮版本，重命名备份文件")

if __name__ == "__main__":
    main()
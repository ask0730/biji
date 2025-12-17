import time
import re
from datetime import datetime, timedelta
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import TimeoutException, NoSuchElementException

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ 未安装python-docx库，Word保存功能将不可用。请运行: pip install python-docx")

class BJSKScraper:
    def __init__(self):
        self.driver = None
        self.articles_data = []
    
    def parse_date(self, date_str):
        """解析日期字符串，返回datetime对象"""
        if not date_str:
            return None
        
        try:
            # 尝试多种日期格式（北京社科使用 2025.10.31 格式）
            date_formats = [
                '%Y.%m.%d',  # 北京社科格式：2025.10.31
                '%Y-%m-%d',
                '%Y年%m月%d日',
                '%Y/%m/%d',
                '%Y-%m-%d %H:%M:%S',
                '%Y/%m/%d %H:%M:%S',
                '%Y.%m.%d %H:%M:%S'
            ]
            
            for fmt in date_formats:
                try:
                    return datetime.strptime(date_str, fmt)
                except:
                    continue
            
            # 如果都不匹配，尝试正则提取日期部分（支持点分隔）
            date_match = re.search(r'(\d{4})[.\-\/年](\d{1,2})[.\-\/月](\d{1,2})', date_str)
            if date_match:
                year, month, day = date_match.groups()
                return datetime(int(year), int(month), int(day))
            
            return None
        except:
            return None
    
    def filter_articles_by_date_range(self, articles, start_date=None, end_date=None, days=None):
        """根据日期范围过滤文章"""
        if not articles:
            return []
        
        # 如果指定了天数，计算开始日期
        if days is not None:
            start_date = datetime.now() - timedelta(days=days)
            end_date = datetime.now()
        elif start_date is None:
            # 默认使用最近7天
            start_date = datetime.now() - timedelta(days=7)
            end_date = datetime.now()
        
        # 确保start_date和end_date是datetime对象
        if isinstance(start_date, str):
            start_date = self.parse_date(start_date)
        if isinstance(end_date, str):
            end_date = self.parse_date(end_date)
        
        if start_date:
            start_date = start_date.replace(hour=0, minute=0, second=0, microsecond=0)
        if end_date:
            end_date = end_date.replace(hour=23, minute=59, second=59, microsecond=999999)
        
        filtered = []
        
        for article in articles:
            publish_time = article.get('publish_time', '')
            if not publish_time:
                # 如果没有发布时间，尝试从链接中提取日期
                link = article.get('link', '')
                if link:
                    # 尝试从URL中提取日期（北京社科可能使用不同的日期格式）
                    date_match = re.search(r'/(\d{4})[/-](\d{2})[/-](\d{2})/', link)
                    if date_match:
                        year, month, day = date_match.groups()
                        publish_time = f"{year}-{month}-{day}"
                        article['publish_time'] = publish_time
            
            if publish_time:
                article_date = self.parse_date(publish_time)
                if article_date:
                    article_date = article_date.replace(hour=0, minute=0, second=0, microsecond=0)
                    # 检查是否在日期范围内
                    if start_date and article_date < start_date:
                        continue
                    if end_date and article_date > end_date:
                        continue
                    filtered.append(article)
                else:
                    # 如果无法解析日期，默认包含（避免遗漏）
                    print(f"  ⚠️ 文章 '{article.get('title', '')}' 日期格式无法解析，将包含在内")
                    filtered.append(article)
            else:
                # 如果没有日期信息，默认包含（避免遗漏）
                print(f"  ⚠️ 文章 '{article.get('title', '')}' 没有日期信息，将包含在内")
                filtered.append(article)
        
        return filtered
    
    def load_config(self, config_file="config.txt"):
        """从配置文件读取设置"""
        config = {
            'days': 7,  # 默认7天
            'start_date': None,
            'end_date': None,
            'output_dir': '项目申报文章'  # 默认输出目录
        }
        
        try:
            import os
            if os.path.exists(config_file):
                with open(config_file, 'r', encoding='utf-8') as f:
                    for line in f:
                        line = line.strip()
                        # 跳过空行和注释
                        if not line or line.startswith('#'):
                            continue
                        
                        # 解析配置项
                        if '=' in line:
                            key, value = line.split('=', 1)
                            key = key.strip()
                            value = value.strip()
                            
                            if key == 'days':
                                try:
                                    config['days'] = int(value)
                                except:
                                    pass
                            elif key == 'start_date':
                                config['start_date'] = value
                            elif key == 'end_date':
                                config['end_date'] = value
                            elif key == 'output_dir':
                                config['output_dir'] = value
        except Exception as e:
            print(f"⚠️ 读取配置文件失败: {e}，使用默认设置（最近7天）")
        
        return config

    def setup_driver(self):
        """设置Chrome浏览器驱动"""
        try:
            chrome_options = Options()
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--disable-gpu')
            chrome_options.add_argument('--window-size=1920,1080')
            chrome_options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
            
            self.driver = webdriver.Chrome(options=chrome_options)
            self.driver.implicitly_wait(10)
            return True
            
        except Exception as e:
            print(f"浏览器启动失败: {e}")
            return False

    def search_articles_by_keyword(self, keyword="项目申报"):
        """搜索包含关键词的文章"""
        articles = []
        
        try:
            # 等待页面加载
            time.sleep(3)
            
            # 方法1: 查找所有链接，筛选包含关键词的标题
            all_links = self.driver.find_elements(By.TAG_NAME, "a")
            
            print(f"🔍 找到 {len(all_links)} 个链接，正在筛选包含'{keyword}'的文章...")
            
            for link in all_links:
                try:
                    title = link.text.strip()
                    href = link.get_attribute('href')
                    
                    # 检查标题是否包含关键词
                    if title and keyword in title and len(title) > 3:
                        # 获取链接的完整URL
                        if href:
                            if href.startswith('/'):
                                # 相对路径，需要拼接
                                current_url = self.driver.current_url
                                base_url = '/'.join(current_url.split('/')[:3])
                                full_url = base_url + href
                            elif href.startswith('http'):
                                full_url = href
                            else:
                                full_url = href
                        else:
                            full_url = ''
                        
                        # 尝试获取发布时间（如果存在）
                        publish_time = ''
                        try:
                            # 查找父元素或兄弟元素中的日期
                            parent = link.find_element(By.XPATH, "./..")
                            parent_text = parent.text
                            
                            # 尝试提取日期格式（优先支持北京社科格式：2025.10.31）
                            date_patterns = [
                                r'(\d{4}\.\d{1,2}\.\d{1,2})',  # 北京社科格式：2025.10.31
                                r'(\d{4}-\d{1,2}-\d{1,2})',
                                r'(\d{4}年\d{1,2}月\d{1,2}日)',
                                r'(\d{4}/\d{1,2}/\d{1,2})'
                            ]
                            
                            for pattern in date_patterns:
                                match = re.search(pattern, parent_text)
                                if match:
                                    publish_time = match.group(1)
                                    break
                        except:
                            pass
                        
                        article = {
                            'title': title,
                            'link': full_url,
                            'publish_time': publish_time
                        }
                        
                        # 避免重复
                        if article not in articles:
                            articles.append(article)
                            print(f"✅ 找到文章: {title}")
                
                except Exception as e:
                    continue
            
            # 方法2: 查找所有包含文本的元素
            try:
                # 使用XPath查找包含关键词的文本节点
                xpath_query = f"//*[contains(text(), '{keyword}')]"
                matching_elements = self.driver.find_elements(By.XPATH, xpath_query)
                
                for element in matching_elements:
                    try:
                        text = element.text.strip()
                        tag_name = element.tag_name.lower()
                        
                        # 如果是链接，直接提取
                        if tag_name == 'a':
                            title = text
                            href = element.get_attribute('href')
                            
                            if title and keyword in title and len(title) > 3:
                                if href:
                                    if href.startswith('/'):
                                        current_url = self.driver.current_url
                                        base_url = '/'.join(current_url.split('/')[:3])
                                        full_url = base_url + href
                                    elif href.startswith('http'):
                                        full_url = href
                                    else:
                                        full_url = href
                                else:
                                    full_url = ''
                                
                                article = {
                                    'title': title,
                                    'link': full_url,
                                    'publish_time': ''
                                }
                                
                                # 避免重复
                                if article not in articles:
                                    articles.append(article)
                                    print(f"✅ 找到文章: {title}")
                        
                        # 如果不是链接，尝试查找父级或子级链接
                        else:
                            try:
                                # 尝试在父元素中查找链接
                                parent = element.find_element(By.XPATH, "./..")
                                parent_links = parent.find_elements(By.TAG_NAME, "a")
                                
                                for parent_link in parent_links:
                                    link_text = parent_link.text.strip()
                                    if keyword in link_text and len(link_text) > 3:
                                        href = parent_link.get_attribute('href')
                                        
                                        if href:
                                            if href.startswith('/'):
                                                current_url = self.driver.current_url
                                                base_url = '/'.join(current_url.split('/')[:3])
                                                full_url = base_url + href
                                            elif href.startswith('http'):
                                                full_url = href
                                            else:
                                                full_url = href
                                        else:
                                            full_url = ''
                                        
                                        article = {
                                            'title': link_text,
                                            'link': full_url,
                                            'publish_time': ''
                                        }
                                        
                                        if article not in articles:
                                            articles.append(article)
                                            print(f"✅ 找到文章: {link_text}")
                            except:
                                pass
                    except:
                        continue
            except Exception as e:
                print(f"XPath搜索时出错: {e}")
            
            return articles
            
        except Exception as e:
            print(f"搜索文章失败: {e}")
            import traceback
            traceback.print_exc()
            return []

    def extract_article_content(self, article_url):
        """访问文章链接并提取文章内容和发布日期"""
        try:
            print(f"  📄 正在访问文章: {article_url}")
            self.driver.get(article_url)
            # 增加等待时间，确保页面完全加载
            time.sleep(5)
            
            content = ""
            publish_date = ""
            article_title = ""  # 从详情页提取标题
            
            # 尝试从详情页提取标题（更准确）
            try:
                # 尝试多种标题选择器
                title_selectors = [
                    "h1",
                    "h2.title",
                    ".title",
                    "[class*='title']",
                    "div.title",
                    "article h1",
                    "article h2"
                ]
                
                for selector in title_selectors:
                    try:
                        title_elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                        for title_elem in title_elements:
                            title_text = title_elem.text.strip()
                            if title_text and len(title_text) > 5:  # 至少5个字符
                                article_title = title_text
                                print(f"  📌 从详情页提取到标题: {article_title}")
                                break
                        if article_title:
                            break
                    except:
                        continue
            except:
                pass
            
            # 尝试从URL中提取日期
            date_match = re.search(r'/(\d{4})[/-](\d{2})[/-](\d{2})/', article_url)
            if date_match:
                year, month, day = date_match.groups()
                publish_date = f"{year}-{month}-{day}"
            
            # 尝试从页面中提取发布日期
            date_selectors = [
                "span.publish-time",
                "div.publish-time",
                "span.date",
                "div.date",
                "time",
                "[class*='date']",
                "[class*='time']",
                "[class*='publish']"
            ]
            
            for selector in date_selectors:
                try:
                    elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    for element in elements:
                        text = element.text.strip()
                        if text and re.search(r'\d{4}', text):
                            parsed_date = self.parse_date(text)
                            if parsed_date:
                                publish_date = parsed_date.strftime("%Y-%m-%d")
                                break
                    if publish_date:
                        break
                except:
                    continue
            
            # 先移除script和style标签，避免干扰
            try:
                self.driver.execute_script("""
                    var scripts = document.getElementsByTagName('script');
                    for(var i = scripts.length - 1; i >= 0; i--) {
                        scripts[i].parentNode.removeChild(scripts[i]);
                    }
                    var styles = document.getElementsByTagName('style');
                    for(var i = styles.length - 1; i >= 0; i--) {
                        styles[i].parentNode.removeChild(styles[i]);
                    }
                    var navs = document.getElementsByTagName('nav');
                    for(var i = navs.length - 1; i >= 0; i--) {
                        navs[i].style.display = 'none';
                    }
                    var footers = document.getElementsByTagName('footer');
                    for(var i = footers.length - 1; i >= 0; i--) {
                        footers[i].style.display = 'none';
                    }
                    var headers = document.querySelectorAll('header, .header, #header');
                    for(var i = 0; i < headers.length; i++) {
                        headers[i].style.display = 'none';
                    }
                """)
                time.sleep(1)
            except:
                pass
            
            # 尝试多种选择器来定位文章内容（北京社科网站）
            content_selectors = [
                "div.content",
                "div.article-content",
                "div.article-body",
                "div.text",
                "div.main-content",
                "div.detail-content",
                "div.news-content",
                "article",
                "div[class*='content']",
                "div[class*='article']",
                "div[class*='detail']",
                "div[class*='text']",
                "#content",
                "#article",
                "#detail",
                ".content",
                ".article",
                ".detail"
            ]
            
            for selector in content_selectors:
                try:
                    elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    if elements:
                        # 选择最长的文本内容（通常是正文）
                        for element in elements:
                            text = element.text.strip()
                            if len(text) > len(content) and len(text) > 200:
                                # 检查是否包含较多中文
                                chinese_chars = len(re.findall(r'[\u4e00-\u9fa5]', text))
                                if chinese_chars > 50:
                                    content = text
                        if content:
                            break
                except:
                    continue
            
            # 如果还没找到，尝试查找包含最多文本的div
            if not content or len(content) < 200:
                try:
                    divs = self.driver.find_elements(By.TAG_NAME, "div")
                    best_div = None
                    best_score = 0
                    
                    for div in divs:
                        try:
                            text = div.text.strip()
                            if len(text) > 500:
                                chinese_chars = len(re.findall(r'[\u4e00-\u9fa5]', text))
                                if '未能提取' not in text and '提取失败' not in text:
                                    score = chinese_chars * 0.7 + len(text) * 0.3
                                    if score > best_score and chinese_chars > 100:
                                        best_score = score
                                        best_div = div
                        except:
                            continue
                    
                    if best_div:
                        content = best_div.text.strip()
                        print(f"  📝 通过div遍历找到内容，长度: {len(content)}")
                except Exception as e:
                    print(f"  ⚠️ div查找失败: {e}")
            
            if content:
                print(f"  ✅ 成功提取内容，长度: {len(content)}字符")
            else:
                print(f"  ⚠️ 未能提取到文章内容")
            
            return content, publish_date, article_title
            
        except Exception as e:
            print(f"  ⚠️ 提取文章内容失败: {e}")
            import traceback
            traceback.print_exc()
            return "", "", ""

    def scrape_articles(self, url, keyword="项目申报", extract_content=True, max_articles=None):
        """抓取包含关键词的文章"""
        if not self.setup_driver():
            return []
        
        try:
            print(f"🌐 正在访问: {url}")
            self.driver.get(url)
            time.sleep(5)  # 等待页面加载
            
            print(f"🔍 正在搜索包含'{keyword}'的文章...")
            articles = self.search_articles_by_keyword(keyword)
            
            # 如果设置了最大文章数，只处理前N篇
            if max_articles and len(articles) > max_articles:
                articles = articles[:max_articles]
                print(f"📌 仅处理前 {max_articles} 篇文章（测试模式）")
            
            # 如果设置了提取内容，访问每个文章链接并提取内容
            if extract_content and articles:
                print(f"\n📖 开始提取 {len(articles)} 篇文章的内容...")
                for i, article in enumerate(articles, 1):
                    if article.get('link'):
                        print(f"\n[{i}/{len(articles)}] 正在提取文章内容...")
                        content, publish_date, extracted_title = self.extract_article_content(article['link'])
                        article['content'] = content
                        # 如果从详情页提取到了标题，更新标题（更准确）
                        if extracted_title:
                            article['title'] = extracted_title
                            print(f"  📌 更新标题: {extracted_title}")
                        # 如果提取到了日期且原来没有日期，则更新
                        if publish_date and not article.get('publish_time'):
                            article['publish_time'] = publish_date
                            print(f"  📅 提取到发布日期: {publish_date}")
                        if content:
                            print(f"  ✅ 成功提取内容（{len(content)} 字符）")
                        else:
                            print(f"  ⚠️ 未能提取到内容")
                        time.sleep(2)  # 避免请求过快
            
            self.articles_data = articles
            return articles
            
        except Exception as e:
            print(f"抓取过程失败: {e}")
            import traceback
            traceback.print_exc()
            return []
        
        # 不在这里关闭浏览器，让调用者决定何时关闭

    def save_article_to_word(self, article, output_dir="."):
        """将单篇文章保存为单独的Word文档"""
        if not DOCX_AVAILABLE:
            print("❌ python-docx库未安装，无法保存Word文档")
            print("   请运行: pip install python-docx")
            return False
        
        if not article:
            return False
        
        try:
            # 创建Word文档
            doc = Document()
            
            # 文章标题
            title = article.get("title", "无标题")
            # 清理标题中的非法字符（Windows文件名不允许的字符）
            # 移除换行符、制表符等空白字符
            safe_title = title.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')
            # 移除Windows文件名不允许的字符
            safe_title = re.sub(r'[<>:"/\\|?*\n\r\t]', '_', safe_title)
            # 移除多余的空格
            safe_title = re.sub(r'\s+', ' ', safe_title).strip()
            if len(safe_title) > 50:  # 限制文件名长度
                safe_title = safe_title[:50]
            
            article_title = doc.add_heading(title, 0)
            article_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加分隔线
            doc.add_paragraph('─' * 50)
            
            # 文章内容
            content = article.get('content', '')
            if content:
                # 将内容按段落分割
                paragraphs = content.split('\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:  # 只添加非空段落
                        para = doc.add_paragraph(para_text)
                        # 设置段落格式
                        para_format = para.paragraph_format
                        para_format.space_after = Pt(6)
                        para_format.first_line_indent = Inches(0.25)  # 首行缩进
            else:
                doc.add_paragraph('（未能提取到文章内容）', style='Intense Quote')
            
            # 文章链接（放在最后）
            if article.get('link'):
                doc.add_paragraph('─' * 50)  # 添加分隔线
                link_para = doc.add_paragraph(f'链接: {article["link"]}')
                # 设置字体大小
                if link_para.runs:
                    link_para.runs[0].font.size = Pt(10)
                link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_title}_{timestamp}.docx"
            # 使用os.path.join确保路径正确（支持Windows和Linux）
            import os
            filepath = os.path.join(output_dir, filename)
            
            # 保存文档
            doc.save(filepath)
            print(f"  ✅ 已保存: {filename}")
            return filepath
            
        except Exception as e:
            print(f"  ❌ 保存文章失败: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def save_articles_to_word(self, articles, output_dir="."):
        """为每篇文章创建单独的Word文档"""
        if not DOCX_AVAILABLE:
            print("❌ python-docx库未安装，无法保存Word文档")
            print("   请运行: pip install python-docx")
            return False
        
        if not articles:
            print("没有数据可保存")
            return []
        
        import os
        # 创建输出目录
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        saved_files = []
        print(f"\n💾 开始保存 {len(articles)} 篇文章到Word文档...")
        
        for i, article in enumerate(articles, 1):
            print(f"\n[{i}/{len(articles)}] 正在保存: {article.get('title', '无标题')}")
            filepath = self.save_article_to_word(article, output_dir)
            if filepath:
                saved_files.append(filepath)
        
        print(f"\n✅ 共保存 {len(saved_files)} 个Word文档到目录: {output_dir}")
        return saved_files

def main():
    print("🚀 开始运行北京社科文章抓取工具...")
    url = "https://www.bjsk.org.cn/newslist-1486-0-0.html"
    keyword = "项目申报"

    try:
        # 读取配置文件
        scraper = BJSKScraper()
        config = scraper.load_config("config.txt")
        
        print("📋 配置信息:")
        if config.get('start_date') and config.get('end_date'):
            print(f"   日期范围: {config['start_date']} 至 {config['end_date']}")
        else:
            print(f"   天数设置: 最近 {config['days']} 天")
        
        print("\n📱 正在启动浏览器...")
        print("🔍 开始抓取文章数据...")
        # 先只搜索文章，不提取内容（节省时间）
        articles = scraper.scrape_articles(url, keyword, extract_content=False)
        
        # 根据配置文件过滤文章（先从URL中提取日期进行初步过滤）
        if articles:
            print(f"\n📅 正在根据配置过滤文章...")
            print(f"   原始文章数: {len(articles)}")
            
            # 先从URL中提取日期，用于初步过滤
            for article in articles:
                if not article.get('publish_time'):
                    link = article.get('link', '')
                    if link:
                        # 尝试多种日期格式
                        date_match = re.search(r'/(\d{4})[/-](\d{2})[/-](\d{2})/', link)
                        if date_match:
                            year, month, day = date_match.groups()
                            article['publish_time'] = f"{year}-{month}-{day}"
            
            # 根据配置进行过滤
            if config.get('start_date') and config.get('end_date'):
                articles = scraper.filter_articles_by_date_range(
                    articles, 
                    start_date=config['start_date'], 
                    end_date=config['end_date']
                )
                print(f"   过滤后文章数: {len(articles)} (日期范围: {config['start_date']} 至 {config['end_date']})")
            else:
                articles = scraper.filter_articles_by_date_range(
                    articles, 
                    days=config['days']
                )
                print(f"   过滤后文章数: {len(articles)} (最近{config['days']}天)")
            
            # 只对过滤后的文章提取内容
            if articles:
                print(f"\n📖 开始提取 {len(articles)} 篇过滤后文章的内容...")
                for i, article in enumerate(articles, 1):
                    if article.get('link'):
                        print(f"\n[{i}/{len(articles)}] 正在提取文章内容...")
                        content, publish_date, extracted_title = scraper.extract_article_content(article['link'])
                        article['content'] = content
                        # 如果从详情页提取到了标题，更新标题（更准确）
                        if extracted_title:
                            article['title'] = extracted_title
                            print(f"  📌 更新标题: {extracted_title}")
                        # 如果提取到了日期且原来没有日期，则更新
                        if publish_date and not article.get('publish_time'):
                            article['publish_time'] = publish_date
                            print(f"  📅 提取到发布日期: {publish_date}")
                        if content:
                            print(f"  ✅ 成功提取内容（{len(content)} 字符）")
                        else:
                            print(f"  ⚠️ 未能提取到内容")
                        time.sleep(2)  # 避免请求过快

        if articles:
            print(f"\n✅ 成功抓取到 {len(articles)} 篇包含'{keyword}'的文章")
            print("\n📋 文章列表:")
            for i, article in enumerate(articles, 1):
                print(f"{i}. {article['title']}")
                if article['link']:
                    print(f"   链接: {article['link']}")
                if article['publish_time']:
                    print(f"   发布时间: {article['publish_time']}")
                if article.get('content'):
                    content_preview = article['content'][:100] + "..." if len(article['content']) > 100 else article['content']
                    print(f"   内容预览: {content_preview}")
                print()
            
            print("\n💾 正在保存到Word文档（每篇文章单独保存）...")
            import os
            output_dir = config.get('output_dir', '项目申报文章')
            print(f"   输出目录: {output_dir}")
            saved_files = scraper.save_articles_to_word(articles, output_dir)
            print(f"\n🎉 抓取完成！共保存 {len(saved_files)} 个Word文档")
        else:
            print(f"❌ 没有找到包含'{keyword}'的文章")
        
        # 所有操作完成后，关闭浏览器
        if scraper.driver:
            print("\n⏳ 浏览器将保持打开5秒，您可以查看结果...")
            time.sleep(5)
            scraper.driver.quit()
            print("✅ 浏览器已关闭")

    except Exception as e:
        print(f"❌ 程序执行失败: {e}")
        import traceback
        traceback.print_exc()
        # 出错时也要关闭浏览器
        if 'scraper' in locals() and scraper.driver:
            scraper.driver.quit()

if __name__ == "__main__":
    main()

